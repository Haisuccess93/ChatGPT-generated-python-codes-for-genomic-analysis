from Bio import SeqIO
from docx import Document

def count_occurrences(sequence, fasta_file):
    count = 0
    with open(fasta_file, "r") as handle:
        for record in SeqIO.parse(handle, "fasta"):
            read_sequence = str(record.seq)
            count += read_sequence.count(sequence)
    return count

def save_to_word_file(output_filename, sequence, occurrences, input_filename):
    doc = Document()
    doc.add_paragraph(f"Sequence searched: {sequence}")
    doc.add_paragraph(f"Occurrences count: {occurrences}")
    doc.add_paragraph(f"Input FASTA file: {input_filename}")
    doc.save(output_filename)

if __name__ == "__main__":
    try:
        sequence = input("Enter the nucleotide sequence to search for: ")
        fasta_file_path = input("Enter the path to the FASTA file: ")
        output_filename = input("Enter the name for the output Word file (without extension): ")

        occurrences = count_occurrences(sequence, fasta_file_path)
        print(f"The number of occurrences in the file is: {occurrences}")

        file_name = fasta_file_path.split("/")[-1].split(".")[0]  # Extracting file name without extension
        save_to_word_file(f"{output_filename}.docx", sequence, occurrences, file_name)
        print(f"Output saved to {output_filename}.docx")
    except FileNotFoundError:
        print("File not found. Please provide a valid file path.")
    except Exception as e:
        print(f"An error occurred: {e}")
